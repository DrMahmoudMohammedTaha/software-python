

import docx

doc = docx.Document("J:/egz.docx")

print(doc)

all_paras = doc.paragraphs
print(len(all_paras))

doc.add_paragraph("This is first paragraph of a MS Word file.")

mydoc.add_heading("This is level 1 heading", 0)
mydoc.add_picture("E:/eiffel-tower.jpg", width=docx.shared.Inches(5), height=docx.shared.Inches(7))
doc.save("J:/egz.docx")

pass

for para in all_paras:
    print(para.text)
    print("-------")


